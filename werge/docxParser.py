from PyPDF2 import PdfFileReader
from tkinter import filedialog
import json
from docx import Document
# parse PDF File contents
import re
from docx.enum.section import WD_ORIENT,WD_SECTION
from os import system
import os
from datetime import datetime
from collections import OrderedDict
from zipfile import BadZipFile
import xml.etree.ElementTree as ET

 

import base64
import configparser
import sys

 

 

BASE_DIR = os.path.dirname(__file__)
build_path = lambda rel_file: os.path.join(BASE_DIR,rel_file)

 

 

config = configparser.ConfigParser()

 

# Will need to be updated when moved to Demisto

config.read_file(open(build_path('config\\config.ini')))#os.path.join(BASE_DIR,'config\\config.ini')))

 

DEBUG = config["DEFAULTS"].getboolean("debug")

 

def exit (error):

    input("Press any Key to Exit")
    if DEBUG:

        raise error

    sys.exit(1)

 

def image_to_b64 ()->base64:

    """
        Load an image and encode it in base64.

        args:
            None
        returns:
            b64_image (base64), img_file_name (str): A tuple of base64 encoded image and the image name taken from the provided file.
            
    """
    try:

        img_file_name = filedialog.askopenfilename(filetypes = [("PNG Image", "*.png"),("JPEG Image","*.jpeg"),("GIF Image","*.gif")])
        
        b64_image = b""

        with open(img_file_name,"rb") as image:
            b64_image = image.read()

        return base64.b64encode(b64_image), img_file_name.split("/")[-1]

    except Exception as e:
        print("Unable to encode image with base64")

        if DEBUG:
            raise e

  

class DocxParagraph:

    """

        A class that handles "paragraphs" found in word docs, converting them into the appropriate json structure for later reproduction by PdfLetter module.

        args:

            required:

            _text (str): The text of the docx.Paragraph object
            _style (docx.Paragraph.style): The docx.style object of the paragraph
            p_type (str), default "paragraph": A string descriptor of the paragraph object
            _paragraph (Paragraph): Instance of docx.Paragraph

           
            Optional:
            font (str), default "<font>{0}</font>": Defines the font tags used for the paragraph text
            paragraph_key (dict): The paragraph key to be used for formated paragraph texts

 

    """

 

    TYPES = ["Paragraph","Spacer","Image","Listed Paragraph","Table"]


    def __init__(self,parser,text,style,p_type:str="Paragraph",paragraph_key="",font:str="<font>{0}</font>",*args,**kwargs):
        self._text = text
        self._style = style
        self.p_type = p_type
        self._paragraph = kwargs.get("paragraph",None)
        self.font = font
        self.content = ""
        self.paragraph_key = paragraph_key
        self._parser = parser

    @property
    def is_empty (self):
        
        return True if re.fullmatch("\s*",self.content) else False


    @property
    def p_type (self):
        return self.__p_type


    @p_type.setter
    def p_type (self,_t):
        if _t in DocxParagraph.TYPES:
            self.__p_type = _t

        else:
            self.__p_type = "Paragraph"

   

    @property
    def is_spacer  (self):
        return True if self._text == "" else False


    @property
    def content_with_font (self):
        return self.font.format(self.content)

    
    
    def prompt_user (self,yes_no_ques:str)->bool:
        """
            Checks if user wants to be prompt and if so prompts the user for some input during the processing of the xml file

            args:
                yes_no_ques (str): The Yes/No question to prompt the user

        """

        while True:
            try:
                if self._parser.prompt:
                    if input(yes_no_ques)[0].lower() == "y":
                        return True
                    else:
                        return False
            except IndexError:
                print("Invalid Entry, Try again")
                
             


    def add_style (self):
        """
            Ensures found word style objects are added to the text_styles with a place holder. Due to the unreliability of parsing word documents,
            placeholders ensure that the style is captured so it can be manually recreated when the JSON structure is reviewed/edited.
        

        """

        new_style =  {"name": self._style.name,  "alignment": "TA_JUSTIFY", "fontName":"arial" }       

 
        if self._style.font.name != None:
            new_style["fontName"] = self._style.font.name

        self._parser.json_structure["text_styles"].append(new_style)

 

    def find_mail_merge_fields (self)->str:
        """

            Checks for mailmerge fields in the paragraph, if found the paragraph content is updated to include the properly formated context_key(s)

            args:
                self
            returns 
                self.content (str): The updated content string


        """


        # set the content to the paragraph text
        # if there are custom variables this will be overridden
        self.content = self._text

        
        # find all MAILMERGE fields, once found ensure that the fieldText is properly converted to a content_key
    
        cust_variables = ""
        
        para_keys = []#["{}".format(datetime.now().timestamp() * 1000)[::4]]
        if self.p_type == "Paragraph" and "MERGEFIELD" in self._paragraph.xml:
           
          
            # I've found a regex search of the paragraph contents is the easiest way of identifiying and extracting the fields
            for a in re.findall(r"(«\S*»)+",self.content):
                merge_field_text = a.replace("»","}").replace("«","{").replace(" ","_")
                merge_field_text = [i[1:] for i in merge_field_text.split("}")[:-1] ] 
                para_keys.extend(merge_field_text)
                
                self._parser.json_structure["data_map"].extend ( merge_field_text  )
                

            self.content = self.content.replace("»","}").replace("«","{").strip()
            
            self.paragraph_key = ":".join(para_keys)
          
            return self.content

        else:

            return self.content


    def check_for_special_styles (self)->str:

        """
            Check the paragraph type, if it's listed or bolded add the appropriate characters and perform an addtional settions/logic relevant to the special type.
            e.g. if an image make sure to add the appropriate json structure, including prompt the user to select the image if set

            args:
                None

    
        """

        self.style = self._parser.json_structure["page_style"]["default_style"]

        namespaces = {"a":"http://schemas.openxmlformats.org/drawingml/2006/main",
                    "w":"http://schemas.openxmlformats.org/wordprocessingml/2006/main",
                    "wp":"http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" }

        if "graphicData" in self._paragraph.xml:
            for a in ET.fromstring(self._paragraph.xml).findall('./w:r/w:drawing//a:graphic//*[@name]',namespaces):
                
                if a.attrib.get("name",False):
                    img_name = a.attrib["name"]
                    self.p_type = "image"
                    self.content = "INSERT IMAGE CONTENT HERE"

                    if self.prompt_user(f"{img_name} was found, load image now? Y/N"):
                        b64_image, self.image_name = image_to_b64()
                        self.content = b64_image.decode('utf-8')
                    
                    else:
                        self.image_name = "untilted graphic"

    
       
        elif self._style.name.lower() == "list paragraph" or re.search(r"^(\u2022|\u2023|\u25E6|\u2043|\u2219).", self.content) != None:
            self.style = "Justify"
            self.font = "<font><bullet>&bull;</bullet>{0}</font>"

        
        
        else:

            # Check if style is already added in the text_style section
            if self._style.name not in [t_style["name"] for t_style in  self._parser.json_structure["text_styles"]]:
                self.add_style()

            self.style = self._style.name

 

        return self.style

 

    @staticmethod
    def convert_to_json (paragraphs,body=True)->dict:

        """

            Provides a dictionary of the paragraphs contents to be added to the page.
            args:
                paragraphs (list): A list of DocxParagraph objects
                body (boolean): Indicates whether the paragraph is part of the body or not
            returns:
                paragraph (generator): A generator of dictionaries representing the contents of the paragraph appropriate for each section, header, body, footer

        """

        for para in paragraphs:
          
            paragraph = {"type":para.p_type,"content":para.content}

            # If the paragraph type is a spacer in the body then return it
            if para.content == "" or para.is_empty:
              
                paragraph["type"] = "Spacer"
                paragraph["spacer"] = para._parser.default_spacer
                paragraph.pop("content",None)
                yield paragraph

           
            # If not s spacer, set the other required keys for a paragraph
            elif para.p_type == "Image":                
                paragraph["name"] = "image_file_location_here"
                paragraph["description"] = ""
                paragraph["hAlign"] = "LEFT"
                paragraph["height"] = 0.4
                paragraph["width"] = 2.0

                yield paragraph

 
            elif para.content != "":
               
                paragraph["style"] = para.style
                paragraph["font"] = para.font
                paragraph["paragraph_key"] = para.paragraph_key
                
            
                yield paragraph

    @classmethod
    def from_word_file (cls,para,parser,para_type="Paragraph"):

        """

            Primary instantiantion for DocxParagraph class.

            args:

                para (json): The json structure that defines the paragraph object
                parser (Parser): The calling parser object
                para_type (str): the paragraph type, used to determines how the DocxParagraph handles itself

 

            returns:

                doc_para (cls): An instance of DocxParagraph

        """


        doc_para = cls(text=para.text,style=para.style,p_type=para_type,paragraph=para._p,parser=parser)
        doc_para.find_mail_merge_fields()
        doc_para.check_for_special_styles()
 

        return doc_para

class DocxImage (DocxParagraph):
    def __init__(self, text:str, style, p_type:str="Table", paragraph_key="", font:str="<font>{0}</font>", *args, **kwargs):
        super().__init__(text, style, p_type=p_type, paragraph_key=paragraph_key, font=font, *args, **kwargs)
        self._table = kwargs["paragraph"]

 

class DocxTable (DocxParagraph):

    def __init__(self, text:str, style, p_type:str="Table", paragraph_key="", font:str="<font>{0}</font>", *args, **kwargs):
        super().__init__(text, style, p_type=p_type, paragraph_key=paragraph_key, font=font, *args, **kwargs)
        self._table = kwargs["paragraph"]


    @property
    def col_headers (self)->dict:

        """

        Provides the column headers for a table. This is the first row of the table       

        """

        col_headers = {}

        for i,c in enumerate(self._table.rows[0].cells):
            col_headers[f"col_{i}"] = {
            "content" : c.text,
            "style" : "Normal",
            "font" : self.font}

 
        return col_headers

 

    @property
    def rows (self)->dict:

        """

            returns the rows with column_number


        """

        row_list = []
        for irow,row in enumerate(self._table.rows):
            row_dict = {f"row_{irow}":[]}
            for icell,c in enumerate(row.cells):
                row_dict[f"row_{irow}"].append(
                    {
                        f"col_num":icell,
                        "style": "Normal",
                        "content":" ".join( [self.find_mail_merge_fields(p) for p in c.paragraphs  ]  ),
                        "font":self.font,
                        "paragraph_key":""

                    }

                )

            row_list.append(row_dict)
        return row_list

   

 

    def find_mail_merge_fields(self,cell_content):

        """

            if there's a paragraph key it'll be set via <<paragraph_key_NAME>> , sets the "font" value.

        """

 

        # set the content to the paragraph text
        # if there are custom variables this will be overridden
        content = cell_content.text



        cust_variables = ""
        # It's simpler and a bit easier to just use a regex search to find the mergefield characters
        # than to parse the xml
        if  "MERGEFIELD" in cell_content._p.xml:
            
            for a in re.findall(r"(«\S*»)+",self.content):
                merge_field_text = a.replace("»","}").replace("«","{").replace(" ","_")
                merge_field_text = [i[1:] for i in merge_field_text.split("}")[:-1] ] 
                para_keys.extend(merge_field_text)
                
                self._parser.json_structure["data_map"].extend ( merge_field_text  )
                

            self.content = self.content.replace("»","}").replace("«","{").strip()
            
            self.paragraph_key = ":".join(para_keys)
          
            return self.content


        else:

            return content

 

    @staticmethod
    def convert_to_json(tables, body) -> dict:
        table_json = []
        for itable,table in enumerate(tables):

            table_json.append(

                    {"table_reference":itable,
                        "headers":table.col_headers,
                    "rows": table.rows
                    }

                )

 

        return  table_json

 

    @classmethod
    def from_word_file(cls, para, parser, para_type):
        table = cls(text="Table",style=para.style,p_type=para_type,paragraph=para,parser=parser)
        table.rows

        return table

 
 

class Parser:

 

    """

        Parser instance handles the loading and parsing of PDF files, converting it's structure into a JSON formated file thats congruent with the expected structure of PdfLetter.py

        args:
      
            file_location (bytes like obj): a file location of a pdf file, must be read using "rb" flag.
            docx_file (Document): The docx.Document instance itself
            paragraphs (Document.paragraphs): The body paragraphs of the document
            header (Document.paragraphs): The headers of the document
            footer (Document.paragraphs): The footers of the document
            base_json_structure (json): A json object used as the base structure that matches the structure used by PdfLetter module. A default one is defined in code but any matching json object can be passed.
            prompt (boolean): Default: false, determines whether the user will be prompted for input regarding how to parse the file or if default values will be used for things like Images.

    """


    def __init__ (self,file_location,docx_file,json_struct,paragraphs=None,header=None,footer=None,tables=None,prompt=False):
        self.file_location = file_location
        self.docx_file = docx_file
        self.paragraphs = paragraphs
        self.header = header
        self.footer = footer
        self.json_structure = json_struct
        self.prompt = prompt
        self.tables = tables
    
    @property
    def default_spacer (self):
        return self.json_structure["default_spacer"]
 

    def build_json (self,file_location=None):

        """

            Creates and saves the json object representing the word document. Uses generators to create the content lists for Headers, Paragraphs and Footer respectively.
            Only supports single pages formats at the moment.

            args:

                file_location (bytes like obj): a file location of a pdf file, must be read using "rb" flag.
                base_json_structure (json): A json object used as the base structure that matches the structure used by PdfLetter module. A default one is defined in code but any matching json object can be passed.
                prompt (boolean): Default: false, determines whether the user will be prompted for input regarding how to parse the file or if default values will be used for things like images.

 

            returns:

                None, saves the created json template

        """


        if file_location == None:

            if config["DEFAULTS"].getboolean("use_docx_file_name"):
                file_location = (self.file_location.name.split("/")[-1]).split(".")[0] + ".json"


            else:

                file_location = input("Save the JSON as:")  


        self.json_structure["pages"]["header"].extend([h  for h in DocxParagraph.convert_to_json(self.header,body=False) ])
        self.json_structure["pages"]["footer"].extend( [f  for f in DocxParagraph.convert_to_json(self.footer,body=False) ])
        self.json_structure["pages"]["tables"].extend([t  for t in DocxTable.convert_to_json(self.tables,body=False) ])
        self.json_structure["pages"]["body"].extend([p for p in DocxParagraph.convert_to_json(self.paragraphs,body=True) ])

       

        with open(file_location,"w") as out:

            json.dump(self.json_structure,fp=out,indent=4)

 
    @property
    def _get_page_layout (self)->dict:

        """Handles just single section layouts.

            Can update to include multiple section layouts later on.

 

            Types of sections:

                https://python-docx.readthedocs.io/en/latest/api/enum/WdSectionStart.html#wdsectionstart

        """

        page_section = self.docx_file.sections[0]

        page_info = {

                    "page_height":page_section.page_height.inches,
                    "page_width":page_section.page_width.inches,
                    "left_margin":page_section.left_margin.inches,
                    "right_margin":page_section.right_margin.inches,
                    "top_margin":page_section.top_margin.inches,
                    "gutter":page_section.gutter.inches

        }

       

 

        return page_info

 

    @staticmethod
    def load_base_json_structure (json_struc_loc)->json:

        """

            read the passed base_json_structure file and return a json object.
            args:
                json_struc_loc (IO): The file location of the base_structure file
            returns:
                json_structure (json): The base_structure file as a json object

        """

 

   

        try:

            json_strut_file_str = ""

           

            # Problably a better  way to specifiy an IO objectin the config file and call it,

            # but I can't seem to find any specific indictation and I don't want to write an ExtendedInterpolation to do it.

            with open (json_struc_loc,"r") as f:

                json_strut_file_str = f.read()

            return json.loads(json_strut_file_str)

 

        except OSError as oe:
            print("OSError when trying to read json_structure file as provided. Make sure the file is not corrupted or in use.")
            exit(oe)

        except json.JSONDecodeError as j:
            print("the provided json_structure is in an invalid format. Make sure the specified format is a valid json item.")
            exit(j)

 

 

    @classmethod

    def from_file (cls,file_location,base_json_structure=build_path(config["DEFAULTS"].get("base_structure_location")),prompt=False):

        """

            import a PDF file from a file.

            arguements:
                file_location (bytes like obj): a file location of a pdf file, must be read using "rb" flag.
                base_json_structure (json): A json object used as the base structure that matches the structure used by PdfLetter module. A default one is defined in config.ini, but any file_location can be passed.
                prompt (boolean): Default: false, determines whether the user will be prompted for input regarding how to parse the file or if default values will be used for things like       Images.

        """ 

        # Read the PDF file contents and create a PDFReader object
        try:

            fileObj = None 
            try:
                fileObj = Document(file_location)
            except BadZipFile as bzp: 
                print("Expected the word file to be a bytes object, got {} instead. Make sure the read mode is set to 'rb'\n\n ".format(type(file_location)))
                exit(bzp)



            # Change to parser instead of word_doc

            word_doc = cls(file_location=file_location,
        
                        docx_file=fileObj,
                        json_struct=Parser.load_base_json_structure(base_json_structure),
                        prompt=prompt

            )

           

            word_doc.paragraphs = [DocxParagraph.from_word_file(doc_para,parser=word_doc)  for doc_para in fileObj.paragraphs]
            word_doc.header = [DocxParagraph.from_word_file(h,para_type="Header",parser=word_doc) for h in fileObj.sections[0].header.paragraphs]
            word_doc.footer = [DocxParagraph.from_word_file(f,para_type="Footer",parser=word_doc) for f in fileObj.sections[0].footer.paragraphs]
            word_doc.tables = [DocxTable.from_word_file(t,para_type="Table",parser=word_doc) for t in fileObj.tables]

            return word_doc

 

 

        except Exception as e:
            print(f"The following error occured when trying to read the file located at {file_location}")
            raise e